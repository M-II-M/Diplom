import whisper
from docxtpl import DocxTemplate
import docx


class TemplFelling:
    def __init__(self, document, checkpoints, audio):
        self.document = document  # Шаблон
        self.checkpoints = checkpoints  # Чекпоинты
        self.audio = audio  # Записанная речь
        self.recognized_text = ""  # Распознанный текст

    def recognize(self):
        """ Распознавание """
        model = whisper.load_model("large")  # Подключаем самую большую модель
        audio = whisper.load_audio(self.audio)

        options = {
            "language": "ru",
            "task": "transcribe",
            "fp16": False  # На моем проце fp16 не поддерживается
        }

        result = whisper.transcribe(model, audio, **options)
        print(result["text"])

        self.recognized_text = result["text"]

        return result["text"]

    def filling_doc(self):
        """ Заполнение """

        checkpoints = self.checkpoints
        text = self.recognized_text

        checkpoints_text = []  # полученный из чекпоинтов текст
        for i in checkpoints:
            start_chekpoint_len = len("Начать " + checkpoints[i].lower() + ".") + 1  # длинна стартового чекпоинта

            start_fill = text.find("Начать " + checkpoints[i].lower() + ".")  # Индекс старта
            end_fill = text.find("Закончить " + checkpoints[i].lower() + ".")  # Индекс конца

            checkpoints_text.append(text[start_fill + start_chekpoint_len:end_fill])

        doc = docx.Document(self.document)
        para_id = []  # Индексы параграфов с чекпоинтами

        # Получение индексов параграфов с чекпоинтами
        for p in range(len(doc.paragraphs)):
            for i in checkpoints:
                if i in doc.paragraphs[p].text:
                    para_id.append(p)

        qwe = 0
        # Добавление чекпоинтов для ввода текста
        for i in range(len(para_id)):
            insert_para = doc.add_paragraph(checkpoints_text[i])
            doc.paragraphs[para_id[i] + qwe]._p.addnext(insert_para._p)
            qwe += 1

        doc.save('test.docx')  # Сохраняем промежуточный документ

        final = DocxTemplate('test.docx')
        context = checkpoints  # Контекст для заполнения чекпоинтов-заголовков

        final.render(context)
        final.save("generated_docx.docx")
